import docx

doc_path = "/Users/user/Downloads/DAMILOLA OGUNALADE-MIT 8333_ CA_301835545_d.ogunlade5545@miva.edu.ng.docx"
doc = docx.Document(doc_path)

# Add Appendix A to the end of the document
doc.add_page_break()
doc.add_heading('Appendix A: Key Source Code Snippets', level=2)

doc.add_heading('1. Authentication Middleware (Backend)', level=3)
doc.add_paragraph('This middleware intercepts incoming requests, verifies the JSON Web Token, and attaches the authenticated user to the request object for role-based access control.')

code_1 = """const jwt = require("jsonwebtoken");
const { User } = require("../models/User");

const authenticateToken = async (req, res, next) => {
  const authHeader = req.headers["authorization"];
  const token = authHeader && authHeader.split(" ")[1];

  if (!token) return res.status(401).json({ error: "Access denied" });

  try {
    const verified = jwt.verify(token, process.env.JWT_SECRET);
    req.user = await User.findById(verified.id);
    next();
  } catch (err) {
    res.status(403).json({ error: "Invalid token" });
  }
};
module.exports = authenticateToken;"""
doc.add_paragraph(code_1)

doc.add_heading('2. Time-Based Greeting Logic (Frontend)', level=3)
doc.add_paragraph('A utility function used across all role dashboards to calculate a personalized greeting based on the local system time.')

code_2 = """export function getGreeting(): string {
  const hour = new Date().getHours();
  if (hour < 12) return "Good morning";
  if (hour < 17) return "Good afternoon";
  return "Good evening";
}"""
doc.add_paragraph(code_2)

doc.save(doc_path)
print("Appendix added.")
