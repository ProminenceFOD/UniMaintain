import docx

doc_path = "/Users/user/Downloads/DAMILOLA OGUNALADE-MIT 8333_ CA_301835545_d.ogunlade5545@miva.edu.ng.docx"
doc = docx.Document(doc_path)

descriptions = {
    "01 Login": "The login interface where users (Students, Staff, Officers, and Admins) authenticate into the system using their email and password. It includes validation and error handling for incorrect credentials.",
    "02 Signup": "The registration page allowing new users to create an account by selecting their role and providing basic details. The UI features a clean, branded sidebar.",
    "Forgot Password": "The interface for users to request a password reset link if they have forgotten their login credentials.",
    "03 Admin Overview Page": "The central dashboard for administrators, displaying high-level system statistics, request distribution charts, and a summary of recent maintenance activity.",
    "04 Admin All Requests Page": "A comprehensive list of all maintenance requests across the university. Admins can filter, search, and assign tasks to specific maintenance officers from this view.",
    "05 Admin User Management Page": "The interface where administrators can view all registered users, toggle their active status, invite new users, and update user roles or details.",
    "06 Admin Analytics Page": "Visualised performance metrics showing request trends over time, resolution rates, and category breakdowns to help administrators make data-driven decisions.",
    "07 Admin Reports Page": "A detailed reporting tool allowing administrators to generate, filter by date range, and export maintenance activity reports in CSV format.",
    "08 Admin Site Settings Page": "Global application settings where administrators can configure system preferences, notification rules, and maintenance categories.",
    "09. API Reference i.e Api Docs Page": "The interactive Swagger API documentation page, which provides a comprehensive reference for all backend endpoints and allows testing of API calls.",
    "10.  Student Dashboard Page": "The landing page for students, providing a quick overview of their active requests and recent notifications, with a welcoming time-based greeting.",
    "11.   Student Requests Page": "A personalized view where students can track the status of all their submitted maintenance requests and filter them by category or status.",
    "12.  Request Details": "A detailed view of a specific maintenance request, displaying the full description, uploaded attachments, and an audit timeline of status updates.",
    "13.  Student Profile Page": "The user profile management screen where students can update their personal information, department, and account preferences.",
    "14 Staff Dashboard": "The dashboard for university staff members, offering a similar overview to students but tailored for staff-specific maintenance tracking.",
    "15 Staff Requests": "The requests management page for staff, allowing them to track and manage issues reported in their offices or respective departments.",
    "16   Staff Profile Page": "The profile settings page for staff members to manage their contact details and system preferences.",
    "Student/ Staff New Maintenance Request Modal": "The submission form where students and staff can create new maintenance requests, complete with category selection, urgency levels, and file uploads.",
    "Student/ Staff  Request Resolved  Modal": "A confirmation modal shown when a maintenance request is successfully marked as resolved.",
    "Student/ Staff  Cancel  this Request confirmation Modal": "A prompt requesting user confirmation before a maintenance request is cancelled to prevent accidental deletions.",
    "17.  Officer Dashboard": "The dashboard for maintenance officers, showing their active workload, pending urgent tasks, and quick access to assigned maintenance jobs.",
    "18.  Officer Assigned Tasks": "A Kanban-style or list view of all maintenance tasks currently assigned to the officer, allowing them to update statuses to 'In Progress' or 'Resolved'.",
    "19 Officer Request Details": "A detailed view for maintenance officers to review the specifics of an assigned job, including location, issue description, and user comments.",
    "Officer Completed": "A historical log of all maintenance requests that the officer has successfully resolved and closed.",
    "20  Officer Profile Page": "The profile page for maintenance officers, displaying their assignment history and account settings.",
    "In-App Notifications": "The real-time notification dropdown alerting users to status changes, new assignments, and important system updates.",
    "Few  Mobile responsiveness  screen": "Examples of the application's responsive design, demonstrating how the interface seamlessly adapts to smaller screens and mobile devices."
}

def insert_paragraph_after(paragraph, text):
    new_p = doc.add_paragraph()
    new_p.text = text
    # Move the new paragraph right after the target paragraph
    p = paragraph._p
    p.addnext(new_p._p)
    return new_p

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for key, desc in descriptions.items():
        if key in text and len(text) < len(key) + 5: # Match mostly the heading exactly
            insert_paragraph_after(p, f"Description: {desc}")
            print(f"Added description for: {key}")

doc.save(doc_path)
print("Done.")
