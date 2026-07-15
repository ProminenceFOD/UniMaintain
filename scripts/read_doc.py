import docx

doc_path = "/Users/user/Downloads/DAMILOLA OGUNALADE-MIT 8333_ CA_301835545_d.ogunlade5545@miva.edu.ng.docx"
doc = docx.Document(doc_path)

for i, p in enumerate(doc.paragraphs[:50]):
    if p.text.strip():
        print(f"[{i}]: {p.text}")

print("--- CHAPTER 4 ---")
for i, p in enumerate(doc.paragraphs):
    if "Implementation" in p.text and "4." in p.text:
        for j in range(i, i+15):
            print(f"[{j}]: {doc.paragraphs[j].text}")
        break

